from PyQt5.QtWidgets import QDialog
from PyQt5 import uic
import docx
import docxtpl
from database import *


class NewTB(QDialog):
    def __init__(self, parent=None):
        self.status_ = True
        self.conf = Ini(self)
        ui_file = self.conf.get_ui("new_TB")
        super(NewTB, self).__init__()
        uic.loadUi(ui_file, self)
        self.parent = parent
        self.table = "workers"
        self.rows_from_db = self.parent.db.get_data("*", self.table)
        if self.rows_from_db == ERR:
            self.status_ = False
            return
        self.b_ok.clicked.connect(self.ev_ok)
        self.b_cancel.clicked.connect(self.ev_cancel)
        self.path = dict()
        self.list_ui = [self.worker_1, self.worker_2, self.worker_3, self.worker_4, self.worker_5,
                        self.worker_6, self.worker_7, self.worker_8, self.worker_9, self.worker_10,
                        self.worker_11, self.worker_12, self.worker_13, self.worker_14, self.worker_15,
                        self.worker_16, self.worker_17, self.worker_18, self.worker_19, self.worker_20,
                        self.worker_21, self.worker_22, self.worker_23, self.worker_24, self.worker_25,
                        self.worker_26, self.worker_27, self.worker_28, self.worker_29, self.worker_30]
        self.list_cb = [NOT for i in range(len(self.list_ui))]
        rows = [self.parent.db.get_data(ALL, WORKERS), self.parent.db.get_data(ALL, ITRS)]
        if ERR in rows:
            return
        self.all_people = rows[0] + rows[1]
        self.init_workers()

    def init_workers(self):
        for item in self.list_ui:
            item.addItem(NOT)
            item.activated[str].connect(self.new_worker)
            item.setEnabled(False)
        self.list_ui[0].setEnabled(True)
        self.all_people.sort()
        workers = list()
        for item in self.all_people:
            if "работает" in item[-2] or "в отпуске" in item[-2]:
                workers.append(item)
        workers.sort(key=lambda x: x[0])
        for people in workers:
            for item in self.list_ui:
                item.addItem(short_name(people))

    def add_item(self, people, ind):
        for item in self.list_ui:
            if self.list_ui.index(item) == ind:
                continue
            man = [item.itemText(i) for i in range(item.count())]
            if not people == NOT:
                man.append(people)
            man.sort()
            i = man.index(people)
            item.insertItem(i, people)
            print([item.itemText(i) for i in range(item.count())])

    def kill_item(self, people, ind):
        for item in self.list_ui:
            if self.list_ui.index(item) == ind:
                continue
            for i in range(item.count()):
                if item.itemText(i) == people:
                    item.removeItem(i)
                    break

    def new_worker(self, some, man=None, indx=None):
        if man:
            val = man
            ind = indx
        else:
            val = self.sender().currentText()
            ind = self.list_ui.index(self.sender())
        flag = True
        if val == NOT:
            if not self.list_cb[ind] == NOT:
                self.add_item(self.list_cb[ind], ind)
                self.list_cb[ind] = NOT
        else:
            self.kill_item(val, ind)
            if self.list_cb[ind] != NOT:
                self.add_item(self.list_cb[ind], ind)
            self.list_cb[ind] = val

        for item in self.list_ui:
            if item.currentText() != NOT:
                item.setEnabled(True)
            else:
                item.setEnabled(flag)
                flag = False

    def get_list_people(self):
        list_people = list()
        for elem in self.list_ui:
            family = elem.currentText()
            if family != NOT:
                item = self.check_row(family)
                if item:
                    list_people.append(item)
        list_people.sort(key=lambda x: x[0])
        return list_people

    def check_row(self, row):
        family = row.split(".")[0][:-2]
        s_name = row.split(".")[0][-1]
        s_sur = row.split(".")[1]
        for item in self.all_people:
            if family == item[0]:
                if s_name == item[1][0]:
                    if s_sur == item[2][0]:
                        return item

    def ev_ok(self):
        people = self.get_list_people()
        self.print_prots(people)
        self.print_study(people)
        self.print_height(people)
        self.create_table(people)
        self.print_cards(people)
        return True

    def ev_cancel(self):
        self.close()

    def create_table(self, list_people):
        list_people = [[]]
        list_people.sort(key=lambda x: x[0][0])
        path = self.conf.get_path("pat_tb") + "/Инструктаж" + DOCX
        doc = docx.Document(path)
        for people in list_people:
            i = next(range(1, len(list_people) + 1))
            doc.tables[0].add_row()
            doc.tables[0].rows[i].cells[0].text = str(i)
            doc.tables[0].rows[i].cells[1].text = full_name(people)  # ФИО
            doc.tables[0].rows[i].cells[2].text = people[4]  # профессия
            doc.tables[0].rows[i].cells[3].text = "Удосторение №" + people[18]
            doc.tables[0].rows[i].cells[4].text = "Договор №" + people[12] + " от " + people[13]
            doc.tables[0].rows[i].cells[5].text = "Выписка из протокола №" + people[17] + " от " + people[19]
            doc.tables[0].rows[i].cells[6].text = people[20] + "/2 от " + people[22]
            doc.tables[0].rows[i].cells[7].text = people[20] + "/1 от " + people[22]
            doc.tables[0].rows[i].cells[8].text = "Допуск на высоту 2 группа (Протокол №" + \
                                                  people[14] + " от " + people[16] + \
                                                  ") Протокол по проверки знаний по электробезопасности " + \
                                                  people[20] + " от " + people[22]
        path = self.conf.get_path("tb") + "/Инструктажи/" + str(dt.datetime.now().date()) + DOCX
        doc.save(path)
        os.startfile(path)

    def print_prots(self, list_people):
        docs = list()
        for people in list_people:
            if not people[20] in docs:
                docs.append(people[20])
                folder = self.conf.get_path("tb_prot") + "/" + people[20]
                for file in os.listdir(folder):
                    os.startfile(folder + "/" + file)
                    ok = msg_info(self, "Открыть следующий документ?")

    def print_study(self, list_people):
        docs = list()
        for people in list_people:
            if not people[17] in docs:
                docs.append(people[17])
                number = people[17].replace("/", "_")
                path = self.conf.get_path("tb_study") + "/" + number + DOCX
                os.startfile(path)
                ok = msg_info(self, "Открыть следующий документ?")
        ok = msg_info(self, "Протоколы на профессий готовы!")

    def print_height(self, list_people):
        docs = list()
        for people in list_people:
            if not people[20] in docs:
                docs.append(people[14])
                path = self.conf.get_path("tb_height") + "/" + people[14] + DOCX
                os.startfile(path)
                ok = msg_info(self, "Открыть следующий документ?")
        ok = msg_info(self, "Протоколы на высоту готовы!")

    def print_cards(self, list_people):
        ok = msg_info(self, "Печатаются личные карточки")
        path = self.conf.get_path("pat_tb") + "/Личная карточка_1" + DOCX
        for i in range(len(list_people)):
            print_to(path)
        ok = msg_info(self, "Положите повторно стопку личных карточек в принтер для печати обратной стороны")
        path = self.conf.get_path("pat_tb") + "/Личная карточка_2" + DOCX
        for i in range(len(list_people)):
            print_to(path)
