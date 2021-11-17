from PyQt5.QtCore import QRegExp as QRE
from PyQt5 import uic
from PyQt5.QtGui import QRegExpValidator as QREVal
from PyQt5.QtWidgets import QDialog
from database import *
from new_template import TempForm
import PyPDF2
import docxtpl
import docx
from my_email import SendPost
msgs = {"mes": "Сообщение", "atn": "Внимание"}


class NewWorker(TempForm):
    def __init__(self, parent=None):
        self.status_ = True
        self.conf = Ini(self)
        self.db = DataBase(self)
        ui_file = self.conf.get_ui("new_worker")
        self.table = WORKERS
        self.rows_from_db = self.db.get_data(ALL, WORKERS)
        super(NewWorker, self).__init__(ui_file, parent)
        # my_pass
        self.cb_vac.activated[str].connect(self.change_vac)
        self.b_send_docs.clicked.connect(self.add_docs)
        self.b_tb.clicked.connect(self.create_tb)
        self.cb_check.stateChanged.connect(self.check_)
        self.init_mask()
        self.list_ui = [self.family, self.name, self.surname, self.post, self.bday,
                        self.passport, self.passport_post,
                        self.adr, self.live_adr, self.phone, self.inn, self.snils,
                        self.n_td, self.d_td, self.n_hght, self.n_card_h,
                        self.d_height, self.n_study, self.n_study_card, self.d_study,
                        self.n_prot, self.n_card, self.d_prot, self.cb_contract,
                        self.d_vac_1, self.d_vac_2, self.place,
                        self.vac_doc, self.cb_vac, self.status]
        self.rows_from_db = self.parent.db.init_list(self.cb_select, ALL, WORKERS, people=True)
        if self.rows_from_db == ERR:
            self.status_ = False
            return
        if self.parent.db.init_list(self.cb_contract, "number, id", CONTRACTS) == ERR:
            self.status_ = False
            return
        self.auto_numbers = ()
        self.my_mem = ""
        self.vac = True
        self.send_docs = True
        self.init_docs()
        self.b_ok.setEnabled(False)

    def check_(self, state):
        self.b_ok.setEnabled(state)

    def add_docs(self):
        if not self.inn.text() or not self.snils.text() or not self.family.text():
            ok = msg_info(self, "Введите для начала ИНН, СНИЛС и ФИО работника")
            if ok:
                return
        ok = msg_info(self, "Отсканируйте все документы нового работника, которые вы хотите отправить бухгалтеру")
        if ok:
            folder = self.parent.conf.get_path("scan")
            files = os.listdir(folder)
            if not files:
                msg_info(self, "Не найдены сканы. Отсканируйте пожалуйста документы")
                return
            check = False
            for file in files:
                if PDF in file:
                    check = True
                    break
            if not check:
                return
            test = os.listdir(self.conf.get_path("workers"))
            if not self.family.text() in test:
                os.mkdir(self.conf.get_path("workers") + "/" + self.family.text())
            path_worker = self.conf.get_path("workers") + "/" + self.family.text() + "/" + "документы" + PDF
            pdf_merger = PyPDF2.PdfFileMerger()
            for doc in files:
                if PDF in doc:
                    pdf_merger.append(str(folder + "/" + doc))
            pdf_merger.write(path_worker)
            os.startfile(path_worker)
            ok = msg_info(self, "Отправить бухгалтеру?")
            if ok:
                if self.inn.text() and self.snils.text():
                    data = "ИНН: " + self.inn.text() + ", СНИЛС: " + self.snils.text()
                    wnd = SendPost(self.parent.db, self.parent.company, path_worker, data)
                    wnd.exec_()

    def init_mask(self):
        symbols = QREVal(QRE("[а-яА-Я ]{30}"))
        number_prot = QREVal(QRE("[А-Яа-я _/- 0-9]{10}"))

        self.family.setValidator(symbols)
        self.name.setValidator(symbols)
        self.surname.setValidator(symbols)
        self.post.setValidator(symbols)
        self.phone.setValidator(QREVal(QRE("[0-9]{11}")))
        self.passport.setValidator(QREVal(QRE("[0-9]{10}")))
        self.inn.setValidator(QREVal(QRE("[0-9]{8}")))
        self.snils.setValidator(QREVal(QRE("[0-9]{8}")))
        self.n_td.setValidator(QREVal(QRE("[0-9]{2}")))
        self.n_hght.setValidator(number_prot)
        self.n_study.setValidator(number_prot)
        self.n_study_card.setValidator(number_prot)
        self.n_prot.setValidator(number_prot)
        self.n_card.setValidator(number_prot)

    def _select(self, text):
        return True

    def get_next_id(self, data):
        g = iter(range(len(self.rows_from_db) + 1))
        for item in self.rows_from_db:
            next(g)
            if data[-1] == item[-1]:
                self.cb_contract.setCurrentIndex(next(g))
                break

    def create_tb(self):
        TB = CreateTB(self)
        TB.exec_()
        ok = msg_info(self, "Протоколы по ОТ, ЭБ и ПТМ успешно созданы!")
        return ok

    def init_docs(self):
        list_docs = [[], [], []]
        for people in self.rows_from_db:
            val_1 = people[20]
            if not val_1 in list_docs[0]:
                list_docs[0].append(val_1)

            val_2 = people[14]
            if not val_2 in list_docs[1]:
                list_docs[1].append(val_2)

            val_3 = people[17]
            if not val_3 in list_docs[2]:
                list_docs[2].append(val_3)
        cb_ = iter([self.cb_docs, self.cb_h, self.cb_st])
        meths = iter([self.select_docs, self.select_h, self.select_st])
        for docs in list_docs:
            docs.sort()
            item = next(cb_)
            item.activated[str].connect(next(meths))
            for doc in docs:
                item.addItem(doc)

    def select_h(self):
        self.select_doc(self.n_hght, self.d_height, 14)

    def select_docs(self):
        self.select_doc(self.n_prot, self.d_prot, 20)

    def select_st(self):
        self.select_doc(self.n_study, self.d_study, 17)

    def select_doc(self, wgt_n, wgt_d, ind):
        doc = self.sender().currentText()
        if doc == NOT:
            return

        wgt_n.setText(doc)
        for people in self.rows_from_db:
            if people[ind] == doc:
                wgt_d.setDate(from_str(people[ind + 2]))
                return


class CreateTB(QDialog):
    def __init__(self, parent=None):
        self.status_ = True
        self.parent = parent
        self.conf = Ini(self)
        ui_file = self.conf.get_ui("new_tb")
        super(CreateTB, self).__init__()
        uic.loadUi(ui_file, self)
        self.list_ui = [self.worker_1, self.worker_2, self.worker_3, self.worker_4, self.worker_5,
                        self.worker_6, self.worker_7, self.worker_8, self.worker_9, self.worker_10,
                        self.worker_11, self.worker_12, self.worker_13, self.worker_14, self.worker_15,
                        self.worker_16, self.worker_17, self.worker_18, self.worker_19, self.worker_20,
                        self.worker_21, self.worker_22, self.worker_23, self.worker_24, self.worker_25,
                        self.worker_26, self.worker_27, self.worker_28, self.worker_29, self.worker_30]
        self.list_cb = [NOT for i in range(len(self.list_ui))]
        self.all_people = self.parent.parent.db.get_data(ALL, WORKERS)
        self.b_ok.clicked.connect(self.ev_ok)
        self.b_cancel.clicked.connect(self.close)
        self.parent = parent
        self.conf = Ini(self)
        self.number = ""
        self.path = ""
        self.count = 0
        self.init_workers()

    def init_workers(self):
        for item in self.list_ui:
            item.addItem(NOT)
            item.activated[str].connect(self.new_worker)
            item.setEnabled(False)
        self.list_ui[0].setEnabled(True)
        self.all_people.sort()
        workers = list()
        for item in self.all_people:
            if "работает" in item[-2] or "в отпуске" in item[-2]:
                workers.append(item)
        workers.sort(key=lambda x: x[0])
        for people in workers:
            for item in self.list_ui:
                item.addItem(short_name(people))

    def add_item(self, people, ind):
        for item in self.list_ui:
            if self.list_ui.index(item) == ind:
                continue
            man = [item.itemText(i) for i in range(item.count())]
            if not people == NOT:
                man.append(people)
            man.sort()
            i = man.index(people)
            item.insertItem(i, people)
            print([item.itemText(i) for i in range(item.count())])

    def kill_item(self, people, ind):
        for item in self.list_ui:
            if self.list_ui.index(item) == ind:
                continue
            for i in range(item.count()):
                if item.itemText(i) == people:
                    item.removeItem(i)
                    break

    def new_worker(self, some, man=None, indx=None):
        if man:
            val = man
            ind = indx
        else:
            val = self.sender().currentText()
            ind = self.list_ui.index(self.sender())
        flag = True
        if val == NOT:
            if not self.list_cb[ind] == NOT:
                self.add_item(self.list_cb[ind], ind)
                self.list_cb[ind] = NOT
        else:
            self.kill_item(val, ind)
            if self.list_cb[ind] != NOT:
                self.add_item(self.list_cb[ind], ind)
            self.list_cb[ind] = val

        for item in self.list_ui:
            if item.currentText() != NOT:
                item.setEnabled(True)
            else:
                item.setEnabled(flag)
                flag = False

    def ev_ok(self):
        people = self.get_list_people()

        dict_docs = self.create_dict(people, 14)
        self.create_height_docs(dict_docs)
        self.create_height_card(people)

        dict_docs = self.create_dict(people, 20)
        self.create_prot_docs(dict_docs)
        self.create_prot_cards(dict_docs)

        dict_docs = self.create_dict(people, 17)
        self.create_study_docs(dict_docs)
        self.close()

    def get_list_people(self):
        list_people = list()
        flag = False
        for elem in self.list_ui:
            family = elem.currentText()
            if family != NOT:
                item = self.check_row(family)
                if item:
                    list_people.append(item)
            else:
                if flag:
                    return list_people
                flag = True
        list_people.sort(key=lambda x: x[0])
        return list_people

    def check_row(self, row):
        family = row.split(".")[0][:-2]
        s_name = row.split(".")[0][-1]
        s_sur = row.split(".")[1]
        for item in self.all_people:
            if family == item[0]:
                if s_name == item[1][0]:
                    if s_sur == item[2][0]:
                        return item

    def get_date_doc(self):
        date = dt.datetime.today()
        week = dt.datetime.now().weekday()
        if 0 < week < 5:
            date = dt.datetime.now().date() - dt.timedelta(days=1)
        if week == 0:
            date = dt.datetime.now().date() - dt.timedelta(days=3)
        if week > 4:
            date = dt.datetime.now().date() - dt.timedelta(days=6-week)
        return str(".".join([str(date)[8:], str(date)[5:7], str(date)[:4]]))

    def create_prot_docs(self, dict_data):
        for number in dict_data.keys():
            self.number = number
            self.data = dict_data[number]
            gen = self.create_prot_table()
            next(gen)
            i = iter(range(1, 4))
            try:
                os.mkdir(self.conf.get_path("tb_prot") + "/" + str(number))
            except:
                pass
            for part in types_docs:
                data = dict()
                data["number"] = str(number) + "/" + str(next(i))  #
                data["date"] = self.get_date_doc()
                path = self.conf.get_path("pat_tb") + types_docs[part]
                doc = docxtpl.DocxTemplate(path)
                doc.render(data)
                self.path = self.conf.get_path("tb_prot") + "/" + str(number) + types_docs[part]
                doc.save(self.path)
                next(gen)
                # os.startfile(self.path)
                pass

    def create_prot_table(self):
        titles = iter(("ОТ-", "ПТМ-", "ЭБ-"))
        yield
        for title in titles:
            g = iter(range(1, len(self.data) + 1))
            doc = docx.Document(self.path)
            for item in self.data:
                i = next(g)
                doc.tables[0].add_row()
                doc.tables[0].rows[i].cells[0].text = str(i)
                doc.tables[0].rows[i].cells[1].text = full_name(item)  # ФИО
                doc.tables[0].rows[i].cells[2].text = item[3]   # профессия
                doc.tables[0].rows[i].cells[3].text = 'ООО "Вертикаль"'
                doc.tables[0].rows[i].cells[4].text = "Сдал №" + title + item[21]
            doc.save(self.path)
            yield

    def create_prot_cards(self, dict_data):
        titles = iter(("ОТ-", "ПТМ-", "ЭБ-"))
        for number in dict_data.keys():
            self.number = number
            gen = self.add_card()
            next(gen)
            self.count = len(dict_data[number])
            for part in types_card.keys():
                title = next(titles)
                flag = True
                for people in dict_data[number]:
                    data = dict()
                    data["family"] = full_name(people)
                    data["post"] = people[3]
                    data["n_doc"] = people[20] + "/" + part
                    data["number"] = title + people[21]
                    data["date"] = people[22]
                    date_end = data["date"][:-1] + str(int(data["date"][-1]) + 1)
                    data["date_end"] = date_end
                    path = self.conf.get_path("pat_tb") + types_card[part]
                    doc = docxtpl.DocxTemplate(path)
                    doc.render(data)
                    if flag:
                        number = str(number).replace("/", "_")
                        path = self.conf.get_path("tb_prot") + "/" + str(number) + types_card[part]
                        flag = False
                        doc.save(path)
                    else:
                        path = os.getcwd() + "/ТБ" + types_card[part]
                        doc.save(path)
                        next(gen)
                        #  os.startfile(path_save)

    def add_card(self):
        card_tmp = [os.getcwd() + "/ТБ" + "/" + item + "_уд" + DOCX for item in ["ОТ", "ПТМ", "ЭБ"]]
        card_res = [self.conf.get_path("tb_prot") + "/" + self.number + "/" + item + "_уд" + DOCX for item in ["ОТ", "ПТМ", "ЭБ"]]
        yield
        for ind in range(3):
            for i in range(self.count):
                source_document = docx.Document(card_tmp[ind])
                target_document = docx.Document(card_res[ind])
                if self.count % 2 == 0 and not self.count == 0:
                    target_document.add_page_break()
                for paragraph in source_document.paragraphs:
                    text = paragraph.text
                    target_document.add_paragraph(text)
                yield

    def find_all(self, number, ind):
        workers = list()
        for people in self.all_people:
            if people[ind] == number:
                workers.append(people)
        return workers

    # ____________ Профессия_____________
    def create_dict(self, list_people, ind):
        numbers = list()
        dict_data = dict()
        for people in list_people:
            if not people[ind] in numbers:
                numbers.append(people[ind])
                dict_data[people[ind]] = self.find_all(people[ind], ind)
        return dict_data

    def create_study_docs(self, dict_data):
        for number in dict_data:
            data = dict()
            data["number"] = number
            data["post"] = dict_data[number][0][3]
            data["date"] = dict_data[number][0][19]
            path = self.conf.get_path("pat_tb") + "/Профессия.docx"
            doc = docxtpl.DocxTemplate(path)
            doc.render(data)
            self.path = self.conf.get_path("tb_study") + "/" + str(number) + DOCX
            doc.save(self.path)
            self.create_study_table(dict_data[number])

    def create_study_table(self, list_people):
        doc = docx.Document(self.path)
        g = iter(range(1, len(list_people) + 1))
        for people in list_people:
            i = next(g)
            doc.tables[0].add_row()
            doc.tables[0].rows[i].cells[0].text = str(i)
            doc.tables[0].rows[i].cells[1].text = full_name(people)  # ФИО
            doc.tables[0].rows[i].cells[2].text = people[3]  # профессия
            doc.tables[0].rows[i].cells[3].text = 'ООО "Вертикаль"'
            doc.tables[0].rows[i].cells[4].text = "Сдал №" + people[21]
        doc.save(self.path)

    # _________ Высота_____________

    def create_height_dict(self, list_people):
        numbers = list()
        dict_data = dict()
        for people in list_people:
            if not people[17] in numbers:
                numbers.append(people[14])
            dict_data[people[14]] = people
        return dict_data

    def create_height_docs(self, dict_data):
        for number in dict_data:
            data = dict()
            data["number"] = number
            data["date"] = dict_data[number][0][20]
            path = self.conf.get_path("pat_tb") + "/Высота.docx"
            doc = docxtpl.DocxTemplate(path)
            doc.render(data)
            self.path = self.conf.get_path("tb_height") +"/" + str(number) + DOCX
            doc.save(self.path)
            self.create_height_table(dict_data[number])

    def create_height_table(self, list_people):
        doc = docx.Document(self.path)
        g = iter(range(6, len(list_people) + 7))
        for people in list_people:
            i = next(g)
            ind = iter(range(3, 9))
            doc.tables[0].add_row()
            a = doc.tables[0].cell(i, 1)
            b = doc.tables[0].cell(i, 2)
            cell = a.merge(b)
            doc.tables[0].rows[i].cells[0].text = str(i - 5)
            cell.text = full_name(people)
            doc.tables[0].rows[i].cells[next(ind)].text = people[3]  # профессия
            doc.tables[0].rows[i].cells[next(ind)].text = 'ООО "Вертикаль"'
            doc.tables[0].rows[i].cells[next(ind)].text = "Сдал №" + people[21]
            doc.tables[0].rows[i].cells[next(ind)].text = "2"  # профессия
            doc.tables[0].rows[i].cells[next(ind)].text = "-"
            doc.tables[0].rows[i].cells[next(ind)].text = "допущен"
        doc.save(self.path)
        os.startfile(self.path)

    def create_height_card(self, list_people):
        for people in list_people:
            data = dict()
            data["card"] = people[15]
            data["family"] = people[0]
            data["name"] = people[1]
            data["surname"] = people[2]
            data["date"] = people[16]
            date_end = data["date"][:-1] + str(int(data["date"][-1]) + 1)
            data["date_end"] = date_end

            path = self.conf.get_path("pat_tb") + "/Высота_уд.docx"
            doc = docxtpl.DocxTemplate(path)
            doc.render(data)
            number = str(people[15]).replace("/", "_")
            self.path = self.conf.get_path("tb_height") + "/Уд_" + number + DOCX
            doc.save(self.path)
            pass
