from PyQt5.QtCore import Qt
from pass_template import TempPass
from database import *
import docx


class ToolsPass(TempPass):
    def __init__(self, parent):
        self.status_ = True
        self.conf = Ini(self)
        ui_file = self.conf.get_ui("pass_tools")
        self.table = TOOLS
        super(ToolsPass, self).__init__(ui_file, parent)
        self.cb_contracts.activated[str].connect(self.change_note)
        self.d_to.setDate(dt.datetime.now().date())
        self.data = {"number": "", "date": "", "contract": ""}
        self.grid = [[self.tool_1, self.code_1, self.count_1, self.si_1],
                     [self.tool_2, self.code_2, self.count_2, self.si_2],
                     [self.tool_3, self.code_3, self.count_3, self.si_3],
                     [self.tool_4, self.code_4, self.count_4, self.si_4],
                     [self.tool_5, self.code_5, self.count_5, self.si_5],
                     [self.tool_6, self.code_6, self.count_6, self.si_6]]
        self.contract = ""
        self.my_text = ""
        self.contracts = []
        self.init_contracts()
        self.init_tools()
        self.main_file += "/Ввоз_инструментов.docx"
        self.change_note()

    # инициализация
    def init_contracts(self):
        self.cb_contracts.addItem(NOT)
        contracts = self.db.get_data("number", CONTRACTS)
        if contracts == ERR or not contracts:
            return ERR
        list(map(self.cb_contracts.addItem, list(chain(*contracts))))

    def init_tools(self):
        rows = self.db.get_data("name", TOOLS)
        add_list = set(chain(*rows))
        for ui in self.grid:
            list(map(ui[0].addItem, add_list))

    def init_select(self):
        rows = self.db.get_data("name, date", TOOLS)
        names = list(map(" | ".join, [*rows]))
        list(map(self.cb_select.addItem, names))
        for row in rows:
            self.cb_select.addItem(" | ".join(row))

    # для заполнения текста
    def _get_data(self):
        if NOT in self.cb_contracts.currentText():
            msg_info(self, "Выберите договор")
            return
        rows = self.parent.db.get_data("number, date, object, type_work, place", CONTRACTS)
        contract = ""
        place = ""
        work = ""
        for row in rows:
            if self.cb_contracts.currentText() in row:
                contract = row
                work = row[-2]
                place = row[-1]
                break
        if "цех" in contract[-1].lower():
            place = contract[-1].replace("цех", "цеха")
        if "ремонт" in contract[-2].lower():
            work = contract[-1].replace("ремонт", "по ремонту ")
            work += 'на объекте: "' + contract[-2] + '" '
        self.data["work"] = work + " " + place
        self.data["contract"] = contract[0] + " от " + contract[1]
        self.data["d_to"] = self.d_to.text()

    # обработчики кнопок
    def _ev_ok(self):
        return True

    def change_note(self):
        pass

    def _create_data(self, path):
        try:
            doc = docx.Document(path)
        except:
            msg_er(self, GET_FILE + path)
            return ERR
        i = 2
        for tool in self.grid:
            if tool[0].currentText() == NOT:
                break
            doc.tables[1].add_row()
            doc.tables[1].rows[i].cells[0].text = str(i)
            doc.tables[1].rows[i].cells[1].text = tool[0].currentText() + " " + tool[1].text()
            doc.tables[1].rows[i].cells[2].text = tool[3].currentText()
            doc.tables[1].rows[i].cells[3].text = str(tool[2].value())
            i += 1
        try:
            doc.save(path)
        except:
            return msg_er(self, GET_FILE + path)
        return True

    def check_input(self):
        for key in self.data.keys():
            if self.data[key] == NOT or self.data[key] == "":
                msg_info(self, FULL_ALL)
                return False
        return True